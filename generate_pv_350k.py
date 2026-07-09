#!/usr/bin/env python3
"""
Generate Property Vista C$350K OTE Job Profile document.
Professional Word document for Mike Daser to present to Lenny (Property Vista).
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Constants ──────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1A, 0x3A, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY_FILL = "F2F2F2"
NAVY_HEX = "1A3A5C"
FONT_NAME = "Calibri"
OUTPUT_PATH = "/home/aiciv/portal_uploads/to-portal/PropertyVista_350K_OTE_JobProfile_Toronto.docx"


def set_cell_shading(cell, color_hex):
    """Apply shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, color=None, size=Pt(10), alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_header_row(table, row_idx, texts, col_widths=None):
    """Format a row as a header row with navy background and white text."""
    row = table.rows[row_idx]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        set_cell_shading(cell, NAVY_HEX)
        set_cell_text(cell, text, bold=True, color=WHITE, size=Pt(10))


def add_data_row(table, row_idx, texts, bold_first=False, shade=False):
    """Format a data row."""
    row = table.rows[row_idx]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        if shade:
            set_cell_shading(cell, LIGHT_GRAY_FILL)
        is_bold = bold_first and i == 0
        set_cell_text(cell, text, bold=is_bold, color=DARK_GRAY, size=Pt(10))


def add_section_heading(doc, text, level=1):
    """Add a navy-colored section heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = NAVY
        run.font.name = FONT_NAME
    return heading


def add_body_paragraph(doc, text, bold=False, italic=False, space_after=Pt(6)):
    """Add a body paragraph with consistent formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = FONT_NAME
        run_b.font.size = Pt(11)
        run_b.font.bold = True
        run_b.font.color.rgb = DARK_GRAY
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(3)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p


def create_table(doc, rows, cols):
    """Create a table with consistent styling."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return table


def add_footer(doc, text):
    """Add footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(8)
        run.font.color.rgb = NAVY


def build_document():
    """Build the full Property Vista C$350K OTE document."""
    doc = Document()

    # ── Page Setup ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Set default font ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Also set heading styles to Calibri + Navy
    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = FONT_NAME
        hs.font.color.rgb = NAVY

    # ── TITLE PAGE ─────────────────────────────────────────────────────
    # Add some spacing before title
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Property Vista: What C$350K OTE Gets You in Toronto")
    run.font.name = FONT_NAME
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    title.paragraph_format.space_after = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "VP of Customer Experience & Technical Support\nRole Profile and Dallas Comparison"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GRAY
    subtitle.paragraph_format.space_after = Pt(30)

    # Divider line
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run("_" * 60)
    run.font.color.rgb = NAVY
    run.font.size = Pt(11)
    divider.paragraph_format.space_after = Pt(30)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_line.add_run("July 2026")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY
    date_line.paragraph_format.space_after = Pt(6)

    prep = doc.add_paragraph()
    prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep.add_run("Prepared by The People Group")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY
    run.font.italic = True

    # Page break
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 1: Context
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "1. Context", level=1)

    add_body_paragraph(
        doc,
        "Property Vista is evaluating the cost of a VP of Customer Experience and "
        "Technical Support role. The current employee, Sheena, operates from Dallas, TX. "
        "Leadership (Lenny) has asked: what would a C$350K OTE package get us if we "
        "hired this role in Toronto instead?"
    )
    add_body_paragraph(
        doc,
        "This report answers that question with a role profile, market data, and a "
        "direct Dallas vs. Toronto comparison."
    )

    # ══════════════════════════════════════════════════════════════════
    # SECTION 2: What C$350K OTE Looks Like in Toronto
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "2. What C$350K OTE Looks Like in Toronto", level=1)

    add_body_paragraph(
        doc,
        "At C$350,000 total on-target earnings in the Toronto SaaS market, you are not "
        "hiring a standard VP of Customer Experience. You are hiring a senior executive "
        "with a significantly expanded mandate. Here is what that compensation level buys:"
    )

    # Compensation Structure Table
    add_section_heading(doc, "C$350K OTE Compensation Structure", level=2)

    table1 = create_table(doc, 4, 3)
    add_header_row(table1, 0, ["Component", "Amount (CAD)", "Notes"])
    add_data_row(table1, 1, [
        "Base Salary",
        "C$240,000 - C$260,000",
        "70-75% of OTE as base (standard for VP-level, non-sales)"
    ], bold_first=True)
    add_data_row(table1, 2, [
        "Performance Bonus",
        "C$90,000 - C$110,000",
        "25-30% bonus target, tied to NRR, CSAT, churn, support SLAs"
    ], bold_first=True, shade=True)
    add_data_row(table1, 3, [
        "Total OTE",
        "C$350,000",
        "On-target earnings (base + bonus at 100% achievement)"
    ], bold_first=True)

    doc.add_paragraph()  # spacing

    add_body_paragraph(
        doc,
        "This puts the role at the 85th-90th percentile for VP Customer Experience in "
        "Toronto. For context:"
    )

    add_bullet(doc, "50th percentile VP CX in Toronto: ", bold_prefix="")
    # Replace last bullet with formatted version
    doc.paragraphs[-1].clear()
    p = doc.paragraphs[-1]
    p.style = doc.styles["List Bullet"]
    run_b = p.add_run("50th percentile ")
    run_b.font.name = FONT_NAME
    run_b.font.size = Pt(11)
    run_b.font.color.rgb = DARK_GRAY
    run = p.add_run("VP CX in Toronto: C$185K-C$195K base (C$215K-C$245K OTE)")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY

    add_bullet(doc, "VP CX in Toronto: C$215K-C$230K base (C$270K-C$300K OTE)", bold_prefix="75th percentile ")
    add_bullet(doc, "85th-90th percentile", bold_prefix="The C$350K OTE level: ")

    doc.add_paragraph()
    add_body_paragraph(
        doc,
        "At this level, you attract candidates from companies like Shopify, Wealthsimple, "
        "Clio, or Lightspeed -- not from 40-person startups.",
        italic=True
    )

    # ══════════════════════════════════════════════════════════════════
    # SECTION 3: The Role Profile at C$350K OTE
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "3. The Role Profile at C$350K OTE", level=1)

    h2 = add_section_heading(doc, "VP/SVP, Customer Experience and Technical Support", level=2)
    add_body_paragraph(
        doc,
        "What this person looks like at this price point",
        italic=True,
        space_after=Pt(12)
    )

    # Candidate Profile Table
    table2 = create_table(doc, 11, 2)

    # Set column widths
    for row in table2.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)

    add_header_row(table2, 0, ["Attribute", "At C$350K OTE You Get"])
    add_data_row(table2, 1, [
        "Experience",
        "15+ years in customer-facing leadership, 7+ years at VP/SVP level"
    ], bold_first=True)
    add_data_row(table2, 2, [
        "Team Size",
        "Has managed 40-100+ person CX/Support organizations"
    ], bold_first=True, shade=True)
    add_data_row(table2, 3, [
        "Revenue Impact",
        "Has owned $20M-$100M+ in customer revenue (renewals, expansion, NRR)"
    ], bold_first=True)
    add_data_row(table2, 4, [
        "Scope",
        "Full P&L accountability for CX, Technical Support, possibly Professional Services"
    ], bold_first=True, shade=True)
    add_data_row(table2, 5, [
        "Industry",
        "SaaS/B2B enterprise, ideally PropTech or adjacent vertical (real estate tech, property management software)"
    ], bold_first=True)
    add_data_row(table2, 6, [
        "Education",
        "MBA or equivalent preferred (not required). Strong operators often have non-traditional backgrounds."
    ], bold_first=True, shade=True)
    add_data_row(table2, 7, [
        "Strategic Capability",
        "Board-level reporting, investor updates, customer advisory boards, strategic account management"
    ], bold_first=True)
    add_data_row(table2, 8, [
        "Technical Depth",
        "Can architect support operations (ticketing, knowledge base, SLA frameworks), evaluate and deploy AI/automation"
    ], bold_first=True, shade=True)
    add_data_row(table2, 9, [
        "International",
        "Likely has managed distributed/global teams (NA + international support)"
    ], bold_first=True)
    add_data_row(table2, 10, [
        "Metric Ownership",
        "NRR, GRR, CSAT/NPS, time-to-resolution, first-contact resolution, escalation rate, cost-to-serve"
    ], bold_first=True, shade=True)

    doc.add_paragraph()

    # Responsibilities
    add_section_heading(doc, "Responsibilities at This Level", level=2)

    responsibilities = [
        "Own the full customer lifecycle post-sale: onboarding, adoption, support, renewals, expansion",
        "Lead and scale CX + Technical Support teams (hiring, training, performance management)",
        "Report directly to CEO/COO with board-level visibility",
        "Own Net Revenue Retention (NRR) as the primary success metric",
        "Build and optimize support operations: tiered support model, SLA frameworks, knowledge base, AI/automation",
        "Drive customer health scoring and proactive intervention programs",
        "Manage escalations for strategic accounts",
        "Collaborate with Product on roadmap priorities driven by customer feedback and support data",
        "Establish and lead Customer Advisory Board",
        "Own customer-facing metrics dashboard and present quarterly to leadership",
    ]
    for resp in responsibilities:
        add_bullet(doc, resp)

    # ══════════════════════════════════════════════════════════════════
    # SECTION 4: Dallas vs. Toronto Comparison
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "4. Dallas vs. Toronto Comparison", level=1)

    add_section_heading(doc, "Side-by-Side Cost Analysis", level=2)

    table3 = create_table(doc, 8, 4)

    add_header_row(table3, 0, [
        "Factor", "Dallas, TX (USD)", "Toronto, ON (CAD)", "Toronto (USD Equiv.)"
    ])
    rows_data = [
        ["Base Salary (50th-60th %ile)", "$220,000 - $250,000", "C$185,000 - C$210,000", "$131,000 - $149,000"],
        ["Performance Bonus (20-25%)", "$44,000 - $62,500", "C$37,000 - C$52,500", "$26,000 - $37,000"],
        ["Total OTE (50th-60th)", "$264,000 - $312,500", "C$222,000 - C$262,500", "$157,000 - $186,000"],
        ["Total OTE at C$350K Target", "N/A", "C$350,000", "$248,000"],
        ["Benefits (annual est.)", "$15,000 - $25,000", "C$8,000 - C$12,000", "$5,700 - $8,500"],
        ["Payroll Tax (employer portion)", "~7.65% (FICA)", "~7.5% (CPP + EI)", "Roughly equivalent"],
        ["Total Employer Cost (OTE + benefits + tax)", "$285,000 - $340,000", "C$240,000 - C$290,000\n(at 50th-60th)", "$170,000 - $206,000"],
    ]
    for i, row_data in enumerate(rows_data):
        shade = (i % 2 == 1)
        add_data_row(table3, i + 1, row_data, bold_first=True, shade=shade)

    doc.add_paragraph()
    add_body_paragraph(
        doc,
        "Conversion rate: 1 USD = 1.41 CAD (June 2026)",
        italic=True,
        space_after=Pt(12)
    )

    # The Math for Lenny
    add_section_heading(doc, "The Math for Lenny", level=2)

    add_body_paragraph(
        doc,
        "At market rates (50th-60th percentile), the same VP role costs Property Vista "
        "approximately 35% less in Toronto than in Dallas in USD terms. Even at the premium "
        "C$350K OTE level (85th-90th percentile Toronto), the USD-equivalent cost is "
        "approximately $248,000, which is STILL less than the 75th percentile Dallas cost "
        "of $280,000-$315,000.",
        bold=True
    )

    add_body_paragraph(
        doc,
        "Put simply: C$350K OTE in Toronto buys you a top-10% executive for the price "
        "of a 75th percentile hire in Dallas.",
        bold=True,
        italic=True,
        space_after=Pt(16)
    )

    # What C$350K OTE Buys in Each Market
    add_section_heading(doc, "What C$350K OTE Buys in Each Market", level=2)

    table4 = create_table(doc, 6, 3)
    add_header_row(table4, 0, [
        "Metric",
        "Dallas (US$350K OTE = C$494K)",
        "Toronto (C$350K OTE = US$248K)"
    ])
    market_rows = [
        ["Market Percentile", "75th-85th percentile", "85th-90th percentile"],
        ["Candidate Quality", "Strong VP, good SaaS experience", "Elite VP/SVP, enterprise-grade"],
        ["Team Size Managed", "20-50 people", "40-100+ people"],
        ["Revenue Owned", "$10M-$30M", "$20M-$100M+"],
        ["Risk", "Competitive market, may get poached", "Toronto talent pool deep, slightly lower churn"],
    ]
    for i, row_data in enumerate(market_rows):
        shade = (i % 2 == 1)
        add_data_row(table4, i + 1, row_data, bold_first=True, shade=shade)

    # ══════════════════════════════════════════════════════════════════
    # SECTION 5: Recommendation
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "5. Recommendation", level=1)

    table5 = create_table(doc, 4, 5)
    add_header_row(table5, 0, [
        "Option", "Description", "OTE (CAD)", "OTE (USD Equiv.)", "Market Position"
    ])
    rec_rows = [
        [
            "A: Market Rate Toronto",
            "Hire a strong VP CX at Toronto 50th-60th percentile",
            "C$222K - C$263K",
            "$157K - $186K",
            "Competitive, solid hire"
        ],
        [
            "B: Premium Toronto (C$350K)",
            "Hire an elite VP/SVP at Toronto 85th-90th percentile",
            "C$350K",
            "$248K",
            "Top-tier, overqualified for current company size"
        ],
        [
            "C: Current Dallas Cost",
            "What Sheena's role costs at Dallas market rates",
            "C$373K - C$441K (equiv.)",
            "$264K - $313K",
            "50th-60th percentile Dallas"
        ],
    ]
    for i, row_data in enumerate(rec_rows):
        shade = (i % 2 == 1)
        add_data_row(table5, i + 1, row_data, bold_first=True, shade=shade)

    doc.add_paragraph()

    add_body_paragraph(
        doc,
        "Our recommendation: Option A gets you an excellent VP in Toronto for "
        "C$222K-C$263K OTE. That is 30-40% less than Dallas market rate for the "
        "equivalent role. If Lenny wants to invest at the C$350K level, he is buying "
        "an executive who is meaningfully overqualified for a 40-person company but "
        "positioned to scale the organization to 200+.",
        bold=True
    )

    doc.add_paragraph()

    add_body_paragraph(
        doc,
        "The question for Lenny is not whether C$350K is justified. It is whether "
        "Property Vista is ready for the executive that C$350K attracts.",
        bold=True,
        italic=True
    )

    # ── Footer ─────────────────────────────────────────────────────────
    add_footer(doc, "Prepared by The People Group | Confidential")

    # ── Save ───────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
